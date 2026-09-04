# -*- coding: utf-8 -*-
"""Конвертер ТЗ (markdown, выход траектории) → .docx. Использование:
   python3 utils/tz-md-to-docx.py <input.md> [output.docx]"""
import sys, re, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

inp = sys.argv[1]
out = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(inp)[0] + '.docx'
NAVY = RGBColor(0x1F, 0x3A, 0x5F)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

def add_runs(p, text):
    # split on **bold**
    for i, part in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        r = p.add_run(part)
        if i % 2 == 1:
            r.bold = True

for raw in open(inp, encoding='utf-8').read().split('\n'):
    line = raw.rstrip()
    if not line:
        continue
    if line.startswith('# '):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line[2:]); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    elif line.startswith('## '):
        h = doc.add_heading(line[3:], level=1)
        for r in h.runs: r.font.color.rgb = NAVY
    elif line.startswith('### '):
        h = doc.add_heading(line[4:], level=2)
        for r in h.runs: r.font.color.rgb = NAVY
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet'); add_runs(p, line[2:])
    elif line.startswith('---'):
        doc.add_paragraph('—' * 20).alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('*') and line.endswith('*'):
        p = doc.add_paragraph(); r = p.add_run(line.strip('*')); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
    else:
        p = doc.add_paragraph(); add_runs(p, line)

doc.save(out)
print('saved:', out)
